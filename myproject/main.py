from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_question_bank_docx():
    # 创建文档
    doc = Document()
    
    # 设置标题
    heading = doc.add_heading('2024年全国职业院校技能大赛高职组（信息安全管理与评估赛项）考试题库', 0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 题库数据 (这里放你复制的题目文本，用换行符分隔)
    # 你可以将上面提取的文本全部放在这里
    questions_text = """
1. 关于并行数据库，下列说法错误的是（）。
A、层次结构可以分为两层，顶层是无共享结构，底层是共享内存或共享磁盘结构
B、无共享结构通过最小化共享资源来降低资源竞争，因此具有很高的可扩展性，适合于 OLTP 应用
C、并行数据库系统经常通过负载均衡的方法来提高数据库系统的业务吞吐率
D、并行数据库系统的主要目的是实现场地自治和数据全局透明共享
答案：D

2. 以下不属于入侵监测系统的是（）。
A、AFID 系统
B、SNORT 系统
C、IETF 系统
D、NETEYE 系统
答案：C

(在此处粘贴更多题目...)
    """
    
    # 处理文本并写入
    lines = questions_text.strip().split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        p = doc.add_paragraph()
        run = p.add_run(line)
        
        # 简单的格式化逻辑
        if line[0].isdigit(): # 题目
            run.bold = True
            run.font.size = Pt(12)
        elif line.startswith("答案："): # 答案
            run.font.color.rgb = RGBColor(255, 0, 0) # 红色显示答案
            run.bold = True
        else: # 选项
            run.font.size = Pt(10.5)

    # 保存文件
    file_name = '信息安全赛项题库.docx'
    doc.save(file_name)
    print(f"文件已生成: {file_name}")

if __name__ == "__main__":
    create_question_bank_docx()